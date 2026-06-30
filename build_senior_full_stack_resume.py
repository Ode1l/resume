from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Jiaheng_Li_Senior_Full_Stack_Resume.docx")

NAVY = RGBColor(23, 55, 78)
TEAL = RGBColor(21, 101, 110)
INK = RGBColor(25, 32, 40)
MUTED = RGBColor(82, 96, 109)
RULE = "A9B8C2"
LINK = RGBColor(20, 88, 138)
FONT = "Calibri"


def set_cell_or_run_font(run, size=9.9, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_repeatable_paragraph_style(style, size, color, before=0, after=0, line=1.08, bold=False):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_border(paragraph, color=RULE, size="8", space="2"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_hyperlink(paragraph, text, url, size=9.3):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "14588A")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(round(size * 2)))
    r_pr.extend([r_fonts, color, underline, sz])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Jiaheng Li  |  ")
    set_cell_or_run_font(run, size=8.2, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def create_bullet_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_num_id = max(
        [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))] + [0]
    ) + 1
    num_id = max([int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))] + [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_num_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "40")
    spacing.set(qn("w:line"), "260")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "15656E")
    r_pr.extend([fonts, color])

    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_id = OxmlElement("w:abstractNumId")
    abstract_id.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_id)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    p_pr.append(num_pr)


def keep_together(paragraph, keep_next=False):
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep_next


def add_section_heading(document, text):
    paragraph = document.add_paragraph(style="Resume Section")
    run = paragraph.add_run(text.upper())
    set_cell_or_run_font(run, size=11.7, bold=True, color=NAVY)
    add_bottom_border(paragraph)
    keep_together(paragraph, keep_next=True)
    return paragraph


def add_role_heading(document, company, role, dates, location):
    paragraph = document.add_paragraph(style="Role Heading")
    company_run = paragraph.add_run(company)
    set_cell_or_run_font(company_run, size=9.9, bold=True, color=INK)
    sep = paragraph.add_run("  |  ")
    set_cell_or_run_font(sep, size=9.6, color=MUTED)
    role_run = paragraph.add_run(role)
    set_cell_or_run_font(role_run, size=9.7, bold=True, color=TEAL)
    keep_together(paragraph, keep_next=True)

    meta = document.add_paragraph(style="Role Meta")
    run = meta.add_run(f"{location}  |  {dates}")
    set_cell_or_run_font(run, size=8.8, italic=True, color=MUTED)
    keep_together(meta, keep_next=True)


def add_bullet(document, num_id, text, lead=None):
    paragraph = document.add_paragraph(style="Resume Bullet")
    apply_bullet(paragraph, num_id)
    if lead and text.startswith(lead):
        first = paragraph.add_run(lead)
        set_cell_or_run_font(first, size=9.7, bold=True, color=INK)
        rest = paragraph.add_run(text[len(lead):])
        set_cell_or_run_font(rest, size=9.7, color=INK)
    else:
        run = paragraph.add_run(text)
        set_cell_or_run_font(run, size=9.7, color=INK)
    keep_together(paragraph)
    return paragraph


def add_project(document, name, role, stack, bullets, num_id, links=None):
    heading = document.add_paragraph(style="Project Heading")
    name_run = heading.add_run(name)
    set_cell_or_run_font(name_run, size=9.8, bold=True, color=INK)
    role_run = heading.add_run(f"  |  {role}")
    set_cell_or_run_font(role_run, size=9.4, bold=True, color=TEAL)
    keep_together(heading, keep_next=True)

    meta = document.add_paragraph(style="Project Meta")
    stack_run = meta.add_run(f"Tech: {stack}")
    set_cell_or_run_font(stack_run, size=8.7, italic=True, color=MUTED)
    if links:
        separator = meta.add_run("  |  ")
        set_cell_or_run_font(separator, size=8.7, color=MUTED)
        for index, (label, url) in enumerate(links):
            if index:
                between = meta.add_run("  |  ")
                set_cell_or_run_font(between, size=8.7, color=MUTED)
            add_hyperlink(meta, label, url, size=8.7)
    keep_together(meta, keep_next=True)

    for bullet in bullets:
        add_bullet(document, num_id, bullet)


def build_document():
    document = Document()
    document.core_properties.title = "Jiaheng Li - Senior Full Stack Resume"
    document.core_properties.subject = "Senior Full Stack Software Engineering"
    document.core_properties.author = "Jiaheng Li"
    document.core_properties.keywords = "C#, .NET, Java, Spring Boot, React, TypeScript, SQL, CI/CD, SaaS"

    section = document.sections[0]
    # Named override: NZ resume geometry uses A4 with compact margins.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = document.styles
    set_repeatable_paragraph_style(styles["Normal"], 9.9, INK, after=2.5, line=1.10)

    style_specs = {
        "Resume Section": (11.7, NAVY, 7, 3, 1.0, True),
        "Role Heading": (9.9, INK, 4.5, 0, 1.0, False),
        "Role Meta": (8.8, MUTED, 0, 1, 1.0, False),
        "Resume Bullet": (9.7, INK, 0, 2, 1.10, False),
        "Project Heading": (9.8, INK, 3.5, 0, 1.0, False),
        "Project Meta": (8.7, MUTED, 0, 1, 1.0, False),
        "Compact Body": (9.8, INK, 0, 2, 1.10, False),
    }
    for name, (size, color, before, after, line, bold) in style_specs.items():
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_repeatable_paragraph_style(styles[name], size, color, before, after, line, bold)

    bullet_num_id = create_bullet_numbering(document)

    # ATS-friendly adaptation of the customer_pack header: no layout table.
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name.paragraph_format.space_after = Pt(0)
    run = name.add_run("JIAHENG LI")
    set_cell_or_run_font(run, size=23, bold=True, color=NAVY)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("FULL-STACK SOFTWARE ENGINEER  |  END-TO-END DELIVERY")
    set_cell_or_run_font(run, size=10.8, bold=True, color=TEAL)

    contact = document.add_paragraph()
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(1)
    for text in ["Auckland, New Zealand", "Open to relocate"]:
        if contact.text:
            sep = contact.add_run("  |  ")
            set_cell_or_run_font(sep, size=9.2, color=MUTED)
        item = contact.add_run(text)
        set_cell_or_run_font(item, size=9.2, color=MUTED)
    sep = contact.add_run("  |  ")
    set_cell_or_run_font(sep, size=9.2, color=MUTED)
    add_hyperlink(contact, "odell.lijiaheng@gmail.com", "mailto:odell.lijiaheng@gmail.com", size=9.2)
    sep = contact.add_run("  |  ")
    set_cell_or_run_font(sep, size=9.2, color=MUTED)
    phone = contact.add_run("027 295 9309")
    set_cell_or_run_font(phone, size=9.2, color=MUTED)
    sep = contact.add_run("  |  ")
    set_cell_or_run_font(sep, size=9.2, color=MUTED)
    add_hyperlink(contact, "github.com/Ode1l", "https://github.com/Ode1l", size=9.2)

    visa = document.add_paragraph()
    visa.paragraph_format.space_before = Pt(0)
    visa.paragraph_format.space_after = Pt(3)
    run = visa.add_run("Open Work Visa valid until 25 February 2029  |  Available immediately")
    set_cell_or_run_font(run, size=9.1, bold=True, color=INK)
    add_bottom_border(visa, color=RULE, size="5", space="5")

    add_section_heading(document, "Professional Summary")
    summary = document.add_paragraph(style="Compact Body")
    text = (
        "Full-stack software engineer with around three years of commercial experience across enterprise ERP/SaaS, "
        "engineering software and New Zealand client delivery. Owns the complete lifecycle from stakeholder discovery "
        "and domain modelling through C#/.NET or Java/Spring APIs, React/TypeScript interfaces, relational data design, "
        "testing, CI/CD and release. Proven in converting ambiguous specialist workflows into maintainable systems, "
        "modernising manual processes and improving data-intensive backend reliability."
    )
    run = summary.add_run(text)
    set_cell_or_run_font(run, size=9.8, color=INK)

    add_section_heading(document, "Core Technical Capabilities")
    skills = [
        "Full stack: C#, .NET 9/10, ASP.NET Core Web API, Java, Spring Boot, Spring MVC, Servlets, React 18, Next.js, TypeScript and Vite.",
        "Data: Oracle, MySQL, PostgreSQL, SQLite and MongoDB; SQL Server coursework; relational modelling, normalisation, joins, transactions, indexing and query optimisation; MyBatis, MyBatis-Plus and Prisma ORM.",
        "Architecture: REST API contracts, layered services, domain modelling, SOLID, design patterns, DDD concepts, authentication, concurrent processing, WebRTC, WebSocket and distributed systems fundamentals.",
        "Delivery and cloud: CI/CD, Git/GitHub workflows, Vercel, AWS, Azure, Terraform, Linux, Windows Server, VPS and cloud instance configuration.",
        "AI-enabled engineering: RAG, MCP servers, agent development and repository-level AI coding workflows with controlled Git diffs, testing and review."
    ]
    for item in skills:
        add_bullet(document, bullet_num_id, item)

    add_section_heading(document, "Selected Technical Projects")
    add_project(
        document,
        "Skyline Beam Calculation Platform",
        "Sole Full-stack Developer",
        ".NET 9/10, C#, ASP.NET Core Web API, React, TypeScript, Vite, SQL",
        [
            "Owned discovery through release, guiding structural engineers to identify known and uncertain inputs, outputs and validation rules, then translating an Excel beam-calculation workbook into typed domain models and maintainable services.",
            "Designed the calculation kernel and API contracts for geometry, supports, loads, reactions, shear, moment, design checks and calculation-book output; delivered the browser workflow and database-backed user functions independently."
        ],
        bullet_num_id,
    )
    add_project(
        document,
        "p2p-lockstep-kit",
        "Creator and Open-source Developer",
        "TypeScript, WebRTC DataChannel, WebSocket, finite-state machines",
        [
            "Designed a browser-first lockstep toolkit with signalling, session lifecycle management, state-hash validation, snapshot synchronisation and reconnect recovery for deterministic turn-based applications."
        ],
        bullet_num_id,
        [("GitHub", "https://github.com/Ode1l/p2p-lockstep-kit")],
    )
    add_project(
        document,
        "NX-Cast",
        "Open-source Developer",
        "C, libnx, DLNA/UPnP, SSDP, SOAP, GENA, FFmpeg, libmpv, OpenGL",
        [
            "Built a DLNA DMR media receiver for Nintendo Switch homebrew, integrating network discovery and control protocols with native media playback and rendering pipelines."
        ],
        bullet_num_id,
        [("GitHub", "https://github.com/Ode1l/NX-Cast")],
    )
    add_project(
        document,
        "Haven Build Group",
        "Commercial Client Delivery",
        "React 18, TypeScript, Vercel, responsive UI, SEO",
        [
            "Delivered a NZD 2,000 production website for a non-technical New Zealand client, taking responsibility for requirements, content structure, visual direction, deployment and ongoing IT support."
        ],
        bullet_num_id,
        [("Live site", "https://havenbuildgroup.co.nz")],
    )

    document.add_page_break()
    add_section_heading(document, "Professional Experience")
    add_role_heading(
        document,
        "Skyline Consulting Engineers Limited",
        "Part-time Full-stack Software Developer",
        "Oct 2024 - Dec 2025",
        "Auckland, New Zealand",
    )
    skyline_bullets = [
        "Operated as the company's sole software developer, taking ownership of requirements analysis, architecture, implementation, testing, stakeholder validation, CI/CD and release for an engineering calculation platform.",
        "Rebuilt the SP1 Excel beam calculation workflow as a C#/.NET 9 application and later upgraded it to .NET 10, separating the calculation kernel, API boundary, persistence and React/TypeScript user interface.",
        "Established calculation interfaces and extensible domain models for section libraries, load builders, support conditions and generated results while keeping source data separate from runtime and derived calculation state.",
        "Worked directly with non-technical engineers: used diagrams, annotated screenshots, email feedback and focused meetings to resolve uncertain requirements without unnecessarily interrupting client work."
    ]
    for item in skyline_bullets:
        add_bullet(document, bullet_num_id, item)

    add_role_heading(
        document,
        "Yonyou Network Technology Co., Ltd.",
        "Junior Backend Developer (Java)",
        "Mar 2022 - Dec 2023",
        "Zhengzhou, China",
    )
    yonyou_bullets = [
        "Developed backend APIs and secondary SaaS features for a front-end/back-end separated enterprise ERP platform using Java, Spring Boot, Spring MVC, Servlets, RestTemplate, MyBatis and Oracle.",
        "Built a quick-login integration service and delivered the message-push module end to end; introduced thread-pool processing for workloads exceeding 100,000 records to improve throughput and delivery stability.",
        "Contributed to a custom Java data-migration service processing up to approximately one million records per day, moving from handwritten MyBatis SQL toward MyBatis-Plus mappings to standardise team delivery while retaining custom SQL where necessary.",
        "Worked with complex ERP and finance schemas, decomposing oversized retrieval flows into bounded database queries and application-side data composition where this improved maintainability, parallelism and consistency of implementation across developers.",
    ]
    for item in yonyou_bullets:
        add_bullet(document, bullet_num_id, item)

    add_role_heading(
        document,
        "Optimal Drainage & Traffic Services Limited",
        "Part-time IT Support",
        "Jun 2024 - Oct 2024",
        "Auckland, New Zealand",
    )
    for item in [
        "Provisioned and administered computers, user accounts and business application access for a 10-person team.",
        "Collected and organised digital evidence supporting pre-litigation action in an intellectual-property infringement matter."
    ]:
        add_bullet(document, bullet_num_id, item)

    add_section_heading(document, "Leadership and Collaboration")
    for item in [
        "Mentored an intern at Yonyou through production code examples, enterprise data-retention conventions and proactive check-ins; encouraged questions, explained review feedback and created a low-pressure path for improvement.",
        "Supported University of Auckland teammates with different technical backgrounds by identifying knowledge gaps early, introducing concepts such as TCP sliding windows before they became blockers and taking ownership of well-defined work when teammates were overloaded.",
        "When contributing to another developer's work, matched the existing naming, style and architectural patterns so changes remained coherent and maintainable for the original owner and wider team."
    ]:
        add_bullet(document, bullet_num_id, item)

    add_section_heading(document, "Education")
    education = document.add_paragraph(style="Compact Body")
    run = education.add_run("Master of Information Technology")
    set_cell_or_run_font(run, size=9.5, bold=True, color=INK)
    run = education.add_run(" - The University of Auckland, New Zealand  |  Jul 2024 - Nov 2025")
    set_cell_or_run_font(run, size=9.35, color=INK)
    education_2 = document.add_paragraph(style="Compact Body")
    run = education_2.add_run("Bachelor of Network Engineering")
    set_cell_or_run_font(run, size=9.5, bold=True, color=INK)
    run = education_2.add_run(" - Henan University, China  |  Sep 2018 - Jun 2022")
    set_cell_or_run_font(run, size=9.35, color=INK)

    add_section_heading(document, "References")
    refs = document.add_paragraph(style="Compact Body")
    run = refs.add_run("Hongchang Zhang, Director, Skyline Consulting Engineers Limited")
    set_cell_or_run_font(run, size=9.3, bold=True, color=INK)
    run = refs.add_run("  |  029 128 6676  |  ")
    set_cell_or_run_font(run, size=9.1, color=MUTED)
    add_hyperlink(refs, "skyline.consulting.96@gmail.com", "mailto:skyline.consulting.96@gmail.com", size=9.1)
    ref_2 = document.add_paragraph(style="Compact Body")
    run = ref_2.add_run("Francis Tscheliski")
    set_cell_or_run_font(run, size=9.3, bold=True, color=INK)
    run = ref_2.add_run("  |  020 4118 6617  |  Written reference letter available on request")
    set_cell_or_run_font(run, size=9.1, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
